"""Genera la cotización como .docx editable (python-docx).

Servicio puro: recibe datos ya resueltos (sin DB ni request) y retorna los
bytes del .docx. Solo modo desglose (una fila por línea)."""
import io
from decimal import Decimal

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _fmt(simbolo: str, valor) -> str:
    return f"{simbolo} {Decimal(valor or 0):,.2f}"


def build_cotizacion_docx(
    *,
    orden,
    iva,
    total,
    simbolo: str,
    tipo_doc: str,
    fecha_str: str,
    vigencia_dias: int,
) -> bytes:
    doc = Document()

    # Encabezado
    h = doc.add_paragraph()
    r = h.add_run("DASIC Industrial")
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sr = sub.add_run(f"{tipo_doc} · {orden.folio or ''}")
    sr.bold = True
    sr.font.size = Pt(12)

    meta = doc.add_paragraph(f"Fecha: {fecha_str}    Vigencia: {vigencia_dias} días")
    meta.runs[0].font.size = Pt(9)

    # Vendedor
    vend = orden.vendedor
    p = doc.add_paragraph()
    p.add_run("Atiende: ").bold = True
    p.add_run(vend.nombre if vend else "Equipo DASIC")
    if vend and getattr(vend, "email", None):
        p.add_run(f"  ·  {vend.email}")

    # Cliente
    cli = orden.cliente
    pc = doc.add_paragraph()
    pc.add_run("Cliente: ").bold = True
    pc.add_run(cli.nombre_empresa if cli else "—")
    extras = []
    if cli and getattr(cli, "rfc_tax_id", None):
        extras.append(f"RFC: {cli.rfc_tax_id}")
    if cli and getattr(cli, "contacto_nombre", None):
        extras.append(f"Contacto: {cli.contacto_nombre}")
    if cli and getattr(cli, "email", None):
        extras.append(cli.email)
    if cli and getattr(cli, "telefono", None):
        extras.append(cli.telefono)
    if extras:
        doc.add_paragraph("   ".join(extras)).runs[0].font.size = Pt(9)

    # Contacto de la orden (sub-2): "Atención: <nombre>"
    if getattr(orden, "contacto", None):
        atenc = doc.add_paragraph()
        atenc.add_run("Atención: ").bold = True
        atenc.add_run(orden.contacto.nombre)

    doc.add_paragraph("")

    # Tabla de líneas
    tabla = doc.add_table(rows=1, cols=6)
    tabla.style = "Table Grid"
    hdr = tabla.rows[0].cells
    for i, t in enumerate(["#", "SKU", "Descripción", "Cant.", "P. Unit", "Subtotal"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(t)
        run.bold = True
        run.font.size = Pt(9)

    for idx, d in enumerate(orden.detalles, start=1):
        prod = d.producto
        sku = d.sku_libre or ((prod.sku_comercial or prod.sku) if prod else None) or "—"
        desc = d.descripcion_libre or (prod.nombre if prod else None) or "Producto"
        row = tabla.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(sku)

        dcell = row[2]
        dcell.text = ""
        dcell.paragraphs[0].add_run(desc).font.size = Pt(9)
        if d.mostrar_marca and d.marca:
            mr = dcell.add_paragraph().add_run(f"Marca: {d.marca}")
            mr.font.size = Pt(8)
        if d.clave_prod_serv or d.clave_unidad_sat:
            sr2 = dcell.add_paragraph().add_run(
                f"SAT: {d.clave_prod_serv or '—'} · Unidad: {d.clave_unidad_sat or '—'}"
            )
            sr2.font.size = Pt(8)
        if d.observaciones_linea:
            orun = dcell.add_paragraph().add_run(d.observaciones_linea)
            orun.italic = True
            orun.font.size = Pt(8)

        row[3].text = str(d.cantidad)
        row[4].text = _fmt(simbolo, d.precio_unitario)
        row[5].text = _fmt(simbolo, d.subtotal)

    # Totales
    doc.add_paragraph("")
    for etiqueta, val in (("Subtotal", orden.total), ("IVA", iva), ("Total", total)):
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = pt.add_run(f"{etiqueta}: {_fmt(simbolo, val)}")
        if etiqueta == "Total":
            run.bold = True
            run.font.size = Pt(12)

    # Observaciones
    if orden.observaciones:
        doc.add_paragraph("")
        doc.add_paragraph().add_run("Observaciones:").bold = True
        doc.add_paragraph(orden.observaciones).runs[0].font.size = Pt(9)

    # Términos y condiciones
    if orden.terminos_condiciones:
        doc.add_paragraph("")
        doc.add_paragraph().add_run("Términos y condiciones:").bold = True
        doc.add_paragraph(orden.terminos_condiciones).runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_remision_docx(
    *,
    remision,
    simbolo: str,
    fecha_str: str,
) -> bytes:
    """Genera una remisión como .docx editable.

    Columnas de precio (P. Unit / Subtotal) sólo se incluyen cuando
    ``remision.mostrar_precios`` es verdadero; en ese caso se agrega una fila
    de Total al final de la tabla.
    """
    doc = Document()

    # --- Encabezado -----------------------------------------------------------
    h = doc.add_paragraph()
    r = h.add_run("DASIC Industrial")
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sr = sub.add_run(f"REMISIÓN · {remision.folio or ''}")
    sr.bold = True
    sr.font.size = Pt(12)

    meta_text = f"Fecha: {fecha_str}"
    ov = remision.orden_venta
    if ov and getattr(ov, "folio", None):
        meta_text += f"    Ref. cotización: {ov.folio}"
    meta = doc.add_paragraph(meta_text)
    meta.runs[0].font.size = Pt(9)

    # --- Cliente --------------------------------------------------------------
    # Resolución: si hay orden_venta usa su cliente, si no usa cliente directo.
    cli = (ov.cliente if ov else None) or getattr(remision, "cliente", None)
    pc = doc.add_paragraph()
    pc.add_run("Cliente: ").bold = True
    pc.add_run(cli.nombre_empresa if cli else "—")
    extras = []
    if cli and getattr(cli, "rfc_tax_id", None):
        extras.append(f"RFC: {cli.rfc_tax_id}")
    if cli and getattr(cli, "contacto_nombre", None):
        extras.append(f"Contacto: {cli.contacto_nombre}")
    if cli and getattr(cli, "email", None):
        extras.append(cli.email)
    if cli and getattr(cli, "telefono", None):
        extras.append(cli.telefono)
    if extras:
        doc.add_paragraph("   ".join(extras)).runs[0].font.size = Pt(9)

    # --- Logística ------------------------------------------------------------
    log_p = doc.add_paragraph()
    log_p.add_run("Transportista: ").bold = True
    log_p.add_run(remision.transportista or "—")
    rec_p = doc.add_paragraph()
    rec_p.add_run("Recibido por: ").bold = True
    rec_p.add_run(remision.recibido_por or "—")

    doc.add_paragraph("")

    # --- Tabla de líneas ------------------------------------------------------
    mostrar = bool(remision.mostrar_precios)
    num_cols = 6 if mostrar else 4  # #, SKU, Descripción, Cantidad [, P.Unit, Subtotal]
    tabla = doc.add_table(rows=1, cols=num_cols)
    tabla.style = "Table Grid"
    hdr = tabla.rows[0].cells
    headers = ["#", "SKU", "Descripción", "Cantidad"]
    if mostrar:
        headers += ["P. Unit", "Subtotal"]
    for i, t in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(t)
        run.bold = True
        run.font.size = Pt(9)

    total_remision = Decimal("0")
    for idx, d in enumerate(remision.detalles, start=1):
        row = tabla.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(d.sku or "—")

        dcell = row[2]
        dcell.text = ""
        dcell.paragraphs[0].add_run(d.descripcion or "—").font.size = Pt(9)
        if d.observaciones_linea:
            orun = dcell.add_paragraph().add_run(d.observaciones_linea)
            orun.italic = True
            orun.font.size = Pt(8)

        unidad = d.clave_unidad_sat or "PZA"
        row[3].text = f"{d.cantidad} ({unidad})"

        if mostrar:
            precio = d.precio_unitario or Decimal("0")
            subtotal = d.subtotal or Decimal("0")
            row[4].text = _fmt(simbolo, precio)
            row[5].text = _fmt(simbolo, subtotal)
            total_remision += Decimal(subtotal)

    # --- Total (sólo con precios) ---------------------------------------------
    if mostrar:
        doc.add_paragraph("")
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = pt.add_run(f"Total: {_fmt(simbolo, total_remision)}")
        run.bold = True
        run.font.size = Pt(12)

    # --- Observaciones --------------------------------------------------------
    if remision.observaciones:
        doc.add_paragraph("")
        doc.add_paragraph().add_run("Observaciones:").bold = True
        doc.add_paragraph(remision.observaciones).runs[0].font.size = Pt(9)

    # --- Bloque de firmas -----------------------------------------------------
    doc.add_paragraph("")
    firma_p = doc.add_paragraph()
    firma_p.add_run("_________________________          _________________________")
    firma_p.runs[0].font.size = Pt(9)
    firma_labels = doc.add_paragraph()
    firma_labels.add_run("        Entregó                              Recibió")
    firma_labels.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
